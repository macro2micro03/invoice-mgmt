from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from app import report_docx


def _make_specs():
    return [
        {"spec": "SHD10", "quantity_ton": 3.606},
        {"spec": "SHD13", "quantity_ton": 21.11},
    ]


def _combined_text(doc: Document) -> str:
    paragraph_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    return paragraph_text + "\n" + table_text


def test_generate_material_inspection_docx_contains_header_fields():
    docx_bytes = report_docx.generate_material_inspection_docx(
        project_name="테스트현장 신축공사",
        work_type="건축",
        material_type="철근",
        document_number="건축(자검)-철근-1호",
        sender="김현장",
        receiver="박감리",
        specs=_make_specs(),
        vendor="동경강업(주)/동국제강",
    )
    doc = Document(BytesIO(docx_bytes))
    combined = _combined_text(doc)
    assert "테스트현장 신축공사" in combined
    assert "건축" in combined
    assert "건축(자검)-철근-1호" in combined
    assert "김현장" in combined
    assert "박감리" in combined


def test_generate_material_inspection_docx_contains_material_rows_and_total():
    docx_bytes = report_docx.generate_material_inspection_docx(
        project_name="테스트현장",
        work_type="건축",
        material_type="철근",
        document_number="건축(자검)-철근-1호",
        sender="김현장",
        receiver="박감리",
        specs=_make_specs(),
        vendor="동경강업(주)/동국제강",
    )
    doc = Document(BytesIO(docx_bytes))
    combined = _combined_text(doc)
    assert "SHD10" in combined
    assert "3.606" in combined
    assert "SHD13" in combined
    assert "21.11" in combined
    assert "동경강업(주)/동국제강" in combined
    assert "24.716" in combined  # 3.606 + 21.11 합계


def test_generate_material_inspection_docx_defaults_inspection_result_to_pass():
    docx_bytes = report_docx.generate_material_inspection_docx(
        project_name="테스트현장",
        work_type="건축",
        material_type="철근",
        document_number="건축(자검)-철근-1호",
        sender="김현장",
        receiver="박감리",
        specs=_make_specs(),
        vendor="동경강업(주)/동국제강",
    )
    doc = Document(BytesIO(docx_bytes))
    combined = _combined_text(doc)
    assert "적합" in combined


def test_generate_material_inspection_docx_contains_signature_placeholders():
    docx_bytes = report_docx.generate_material_inspection_docx(
        project_name="테스트현장",
        work_type="건축",
        material_type="철근",
        document_number="건축(자검)-철근-1호",
        sender="김현장",
        receiver="박감리",
        specs=_make_specs(),
        vendor="동경강업(주)/동국제강",
    )
    doc = Document(BytesIO(docx_bytes))
    combined = _combined_text(doc)
    assert combined.count("(인)") >= 2


def test_generate_material_inspection_docx_sets_korean_font_on_every_run():
    # 한글 eastAsia 폰트를 지정하지 않으면 워드가 아닌 일부 뷰어(WPS 등)에서
    # 한글이 네모(□)로 보이는 문제가 있었다. 헤딩/문단/표 셀 모든 run에
    # ascii/eastAsia 폰트가 실제로 박혀 있는지 확인한다.
    docx_bytes = report_docx.generate_material_inspection_docx(
        project_name="테스트현장",
        work_type="건축",
        material_type="철근",
        document_number="건축(자검)-철근-1호",
        sender="김현장",
        receiver="박감리",
        specs=_make_specs(),
        vendor="동경강업(주)/동국제강",
    )
    doc = Document(BytesIO(docx_bytes))

    runs = []
    for paragraph in doc.paragraphs:
        runs.extend(paragraph.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    runs.extend(paragraph.runs)

    assert runs, "문서에 텍스트 run이 하나도 없습니다"
    for run in runs:
        assert run.font.name == report_docx.KOREAN_FONT_NAME
        east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia"))
        assert east_asia == report_docx.KOREAN_FONT_NAME
