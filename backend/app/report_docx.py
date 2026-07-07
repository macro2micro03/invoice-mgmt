from datetime import date
from io import BytesIO

from docx import Document

MATERIAL_HEADERS = ["품명", "규격", "단위", "수량", "반입업체명/제조회사명", "검수결과", "검수자"]


def generate_material_inspection_docx(
    *,
    project_name: str,
    work_type: str,
    material_type: str,
    document_number: str,
    sender: str,
    receiver: str,
    specs: list[dict],
    vendor: str,
) -> bytes:
    today = date.today().strftime("%Y년 %m월 %d일")
    doc = Document()
    doc.add_heading("자재검수요청서/검수결과통보", level=1)

    info_table = doc.add_table(rows=5, cols=4)
    info_table.style = "Table Grid"
    info_rows = info_table.rows
    info_rows[0].cells[0].text = "공 사 명"
    info_rows[0].cells[1].text = project_name
    info_rows[0].cells[2].text = "승인구분"
    info_rows[0].cells[3].text = ""
    info_rows[1].cells[0].text = "공 종 명"
    info_rows[1].cells[1].text = work_type
    info_rows[1].cells[2].text = "문서번호"
    info_rows[1].cells[3].text = document_number
    info_rows[2].cells[0].text = "발 신 자"
    info_rows[2].cells[1].text = sender
    info_rows[2].cells[2].text = "접수일자"
    info_rows[2].cells[3].text = today
    info_rows[3].cells[0].text = "수 신 자"
    info_rows[3].cells[1].text = receiver
    info_rows[3].cells[2].text = "검수일자"
    info_rows[3].cells[3].text = today
    info_rows[4].cells[0].text = "검수위치"
    info_rows[4].cells[1].text = "현장 내"
    info_rows[4].cells[2].text = ""
    info_rows[4].cells[3].text = ""

    doc.add_paragraph("")

    material_table = doc.add_table(rows=1 + len(specs) + 1, cols=len(MATERIAL_HEADERS))
    material_table.style = "Table Grid"
    for cell, header_text in zip(material_table.rows[0].cells, MATERIAL_HEADERS):
        cell.text = header_text

    total_ton = 0.0
    for row_index, spec_row in enumerate(specs, start=1):
        cells = material_table.rows[row_index].cells
        cells[0].text = material_type
        cells[1].text = spec_row["spec"]
        cells[2].text = "Ton"
        cells[3].text = f"{spec_row['quantity_ton']:.3f}".rstrip("0").rstrip(".")
        cells[4].text = vendor
        cells[5].text = "적합"
        cells[6].text = ""
        total_ton += spec_row["quantity_ton"]

    total_cells = material_table.rows[-1].cells
    total_cells[0].text = material_type
    total_cells[1].text = "계"
    total_cells[2].text = "Ton"
    total_cells[3].text = f"{round(total_ton, 3):.3f}".rstrip("0").rstrip(".")

    doc.add_paragraph("")
    doc.add_paragraph(f"위 자재에 대하여 검수를 요청합니다.")
    doc.add_paragraph(f"검수 요청일: {today}    현장 대리인: {sender}")
    doc.add_paragraph(f"위 자재 검수결과를 통보합니다.")
    doc.add_paragraph(f"통보 일자: {today}    총괄 관리원: {receiver}")
    doc.add_paragraph("미승인 사유: ")
    doc.add_paragraph("처리 방안: ")
    doc.add_paragraph("붙임: 1. 반입송장 2. 사진대지")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
